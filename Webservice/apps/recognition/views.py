import datetime
import os
import re
import uuid
import sys
import docx
import speech_recognition as speech_recog
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.http import Http404, HttpResponse
from django.shortcuts import render, redirect
from django.views.decorators.csrf import csrf_exempt
from googletrans import Translator
from gtts import gTTS
from pydub import AudioSegment
from .ysk import sound_to_text
from . import models
from .models import TextAudio

sys.path.append(os.path.abspath('..'))
from texttospeech import models as mod


@login_required
@csrf_exempt
def index(request):
    if (request.method == "POST") and ('Отправить' in request.POST):
        try:
            uploadedFile = request.FILES["uploadedFile"]
            print(type(uploadedFile))
            filename, file_extension = os.path.splitext(uploadedFile.name)
            filename = str(uuid.uuid4()) + file_extension
            uploadedFile.name = filename
            document = models.Document(
                uploadedFile=uploadedFile
            )
            document.save()
            if file_extension == '.mp3':
                sound = AudioSegment.from_mp3('media/Uploaded Files/' + filename)
            else:
                try:
                    sound = AudioSegment.from_file('media/Uploaded Files/' + filename, file_extension.replace('.', ''))
                except:
                    os.remove('media/Uploaded Files/' + filename)
                    messages.info(request, 'Данное расширение файла не поддерживается')
                    return render(request, 'main.html')

            sound.export('media/Result Files/' + filename + '.wav', format="wav")

            sample_audio = speech_recog.AudioFile('media/Result Files/' + filename + '.wav')
            recog = speech_recog.Recognizer()
            with sample_audio as audio_file:
                audio_content = recog.record(audio_file)

            result = recog.recognize_google(audio_content, language='ru-RU')

            os.remove('media/Uploaded Files/' + filename)
            os.remove('media/Result Files/' + filename + '.wav')
            textaudio = models.TextAudio(
                text=result,
                title='Запись от ' + '.'.join(str(datetime.datetime.now()).split('.')[:-1]),
                sender=str(request.user.username)
            )
            textaudio.save()
            messages.info(request,'Запись обработана и добавлена в ваши записи!')
        except:
            messages.info(request, 'Что-то пошло не так :(')
        return render(request, 'main.html')


    elif request.method == "POST":
        uploadedFile = request.FILES["audio_data"]
        filename, file_extension = os.path.splitext(uploadedFile.name)
        filename = str(uuid.uuid4()) + file_extension
        uploadedFile.name = filename
        document = models.Document(
            uploadedFile=uploadedFile
        )
        filenamenew = str(uuid.uuid4()) + '.wav'
        document.save()
        os.rename('media/Uploaded Files/' + filename, 'media/Uploaded Files/' + filenamenew)
        filename = filenamenew
        sample_audio = speech_recog.AudioFile('media/Uploaded Files/' + filename)
        recog = speech_recog.Recognizer()
        with sample_audio as audio_file:
            audio_content = recog.record(audio_file)

        result = recog.recognize_google(audio_content, language='ru-RU')

        os.remove('media/Uploaded Files/' + filename)
        textaudio = models.TextAudio(
            text=result,
            title='Запись от ' + '.'.join(str(datetime.datetime.now()).split('.')[:-1]),
            sender=str(request.user.username)
        )
        textaudio.save()
        return render(request, "main.html")

    else:
        return render(request, "main.html")


@login_required
def history(request):
    latest_text_list = TextAudio.objects.filter(sender=str(request.user.username)).order_by('-id')
    return render(request, 'history.html', {'latest_text_list': latest_text_list})


@csrf_exempt
@login_required
def detail(request, id):
    if request.method == 'GET':
        try:
            a = TextAudio.objects.get(id=id)
        except:
            raise Http404('Запись не найдена!')
        if a.sender == request.user.username:
            return render(request, 'detail.html', {'text': a})
        else:
            return HttpResponse('Запись не найдена')

    elif (request.method == "POST") and ('Отправить запись' in request.POST):
        users = User.objects.all()
        users_list = str(users)
        nameuser = str(request.POST.get('nameuser'))

        if re.search(r'\<User: ' + nameuser + r'\>', users_list) is None or nameuser == str(request.user.username):
            a = TextAudio.objects.get(id=id)
            messages.info(request, 'Такого пользователя не существует или вы пытаетесь отправить себе')
            return render(request, 'detail.html', {'text': a})
        else:
            a = TextAudio.objects.get(id=id)
            a.recipients += str(request.POST.get('nameuser')) + ','
            a.save()
            messages.info(request, 'Запись отправлена пользователю:' + str(request.POST.get('nameuser')))
            return render(request, 'detail.html', {'text': a})

    elif (request.method == "POST") and ('Перевести в аудио' in request.POST):
        a = TextAudio.objects.get(id=id)
        uploadedtext = a.text
        s = gTTS(uploadedtext, lang='ru')
        filename = str(uuid.uuid4()) + '.mp3'
        s.save('media/Uploaded Files/' + filename)
        path = 'media/Uploaded Files/' + filename
        text = mod.AudioText(
            title='Запись от ' + '.'.join(str(datetime.datetime.now()).split('.')[:-1]),
            uploadedFile=path,
            sender=str(request.user.username),
            namefile=filename
        )
        text.save()
        messages.info(request, 'Запись переведена в аудио!')
        return render(request, 'detail.html', {'text': a})

    elif (request.method == "POST") and ('Изменить' in request.POST):
        a = TextAudio.objects.get(id=id)
        text = request.POST.get('refactor')
        a.text = text
        a.save()
        return render(request, 'detail.html', {'text': a})

    elif (request.method == "POST") and ('Удалить' in request.POST):
        a = TextAudio.objects.get(id=id)
        a.delete()
        return redirect('/recognition/history')

    elif (request.method == "POST") and ('Сохранить в Word' in request.POST):
        a = TextAudio.objects.get(id=id)
        doc = docx.Document()
        doc.add_paragraph(a.text)
        doc.save('media/Result Files/' + a.title + '_' + str(a.id) + '.docx')
        return redirect('/media/Result Files/' + a.title + '_' + str(a.id) + '.docx')

    elif (request.method == "POST") and ('Перевести' in request.POST):
        a = TextAudio.objects.get(id=id)
        translator = Translator()
        if request.POST.get('Selectedeng') == '1':
            result = translator.translate(a.text, src='ru', dest='en')
            a.text = result.text
            a.save()
        else:
            result = translator.translate(a.text, src='en', dest='ru')
            a.text = result.text
            a.save()
        return render(request, 'detail.html', {'text': a})


    else:
        return HttpResponse(200)


@login_required
def sendhistory(request):
    user = request.user.username
    latest_text_list = TextAudio.objects.filter(recipients__contains=str(user)).order_by('-id')
    return render(request, 'sendhistory.html', {'latest_text_list': latest_text_list})


@csrf_exempt
@login_required
def detailsend(request, id):
    if request.method == 'GET':
        try:
            a = TextAudio.objects.get(id=id)
        except:
            raise Http404('Запись не найдена!')
        if request.user.username in a.recipients.split(','):
            return render(request, 'senddetail.html', {'text': a})
        else:
            return HttpResponse('Нет доступа к записи!')

    elif (request.method == "POST") and ('Перевести в аудио' in request.POST):
        a = TextAudio.objects.get(id=id)
        uploadedtext = a.text
        s = gTTS(uploadedtext, lang='ru')
        filename = str(uuid.uuid4()) + '.mp3'
        s.save('media/Uploaded Files/' + filename)
        path = 'media/Uploaded Files/' + filename
        text = mod.AudioText(
            title='Запись от ' + '.'.join(str(datetime.datetime.now()).split('.')[:-1]),
            uploadedFile=path,
            sender=str(request.user.username),
            namefile=filename
        )
        text.save()
        messages.info(request, 'Запись переведена в аудио!')
        return render(request, 'senddetail.html', {'text': a})


    elif (request.method == "POST") and ('Сохранить в Word' in request.POST):
        a = TextAudio.objects.get(id=id)
        doc = docx.Document()
        doc.add_paragraph(a.text)
        doc.save('media/Result Files/' + a.title + '_' + str(a.id) + '.docx')
        return redirect('/media/Result Files/' + a.title + '_' + str(a.id) + '.docx')

    elif (request.method == "POST") and ('Перевести' in request.POST):
        a = TextAudio.objects.get(id=id)
        translator = Translator()
        if request.POST.get('Selectedeng') == '1':
            result = translator.translate(a.text, src='ru', dest='en')
            a.text = result.text
        else:
            result = translator.translate(a.text, src='en', dest='ru')
            a.text = result.text
        return render(request, 'senddetail.html', {'text': a})

    elif (request.method == "POST") and ('Удалить' in request.POST):
        a = TextAudio.objects.get(id=id)
        recip = a.recipients
        rec = recip.replace(str(request.user.username) + ',', ' ')
        a.recipients = rec
        a.save()
        return redirect('/recognition/sendhistory')

    else:
        return HttpResponse('200')

@login_required
@csrf_exempt
def yandexkit(request):
    if (request.method == "POST") and ('Отправить' in request.POST):

        iam_token = 't1.9euelZqRy5uayY6Oz5TJloqUzpqWx-3rnpWaiZSNmJmMjJKVzZiTkZSUz43l8_c2XF9t-e8IIzlE_t3z93YKXW357wgjOUT-.rR6gN3w-ogyhcN3cFfLYAY395Obo2Lmts4n63MRF3-hvWJjEcD7CqdOXJFt0TSfhijEQZyaiNWCHtGgG1l7bDQ'
        folder_id = 'b1gcj077528pofd7e90e'

        uploadedFile = request.FILES["uploadedFile"]
        print(type(uploadedFile))
        filename, file_extension = os.path.splitext(uploadedFile.name)
        filename = str(uuid.uuid4()) + file_extension
        uploadedFile.name = filename
        document = models.Document(
            uploadedFile=uploadedFile
        )
        document.save()
        if file_extension == '.mp3':
            sound = AudioSegment.from_mp3('media/Uploaded Files/' + filename)
        else:
            try:
                sound = AudioSegment.from_file('media/Uploaded Files/' + filename, file_extension.replace('.', ''))
            except:
                os.remove('media/Uploaded Files/' + filename)
                messages.info(request, 'Данное расширение файла не поддерживается')
                return render(request, 'yandexkit.html')
        try:
            sound.export('media/Result Files/' + filename + '.ogg', format="ogg")

            with open('media/Result Files/' + filename + '.ogg', "rb") as f:
                data = f.read()

            os.remove('media/Uploaded Files/' + filename)
            os.remove('media/Result Files/' + filename + '.ogg')

            result = sound_to_text(iam_token,folder_id,data)
            textaudio = models.TextAudio(
                text=result,
                title='Запись от ' + '.'.join(str(datetime.datetime.now()).split('.')[:-1]),
                sender=str(request.user.username)
            )
            textaudio.save()
            messages.info(request, 'Запись обработана и добавлена в ваши записи!')
        except:
            messages.info(request, 'Что-то пошло не так :(')
        return render(request, 'yandexkit.html')
    else:
        return render(request, 'yandexkit.html')
